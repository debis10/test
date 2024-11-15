from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

# Create a new Word document
document = Document()

# Set document title with centered alignment
title = document.add_paragraph("BMS INSTITUTE OF TECHNOLOGY AND MANAGEMENT")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run()
title_run.bold = True
title_run.font.size = Pt(16)

# Subtitle
subtitle = document.add_paragraph("Higher Education Facilitation Center")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
document.add_paragraph("\n")

# Form fields
fields = {
    "Full name:": "Prajwal Shenoy",
    "USN:": "1BY18CS229",
    "Address:": "18th Main 2nd Cross, HAL 2nd Stage, Indiranagar",
    "Mobile number:": "8217493742",
    "Email:": "prajwalshenoy42@gmail.com",
    "Department:": "Computer Science & Engineering",
    "Year of passing:": "2022",
    "Field of interest (MS/Ph.D.), Program (e.g., Data Science, CS, Mechanical):": "MSc in Cybersecurity",
    "Universities of interest:": "Saarland University - Germany\nKTH Royal Institute of Technology - Sweden\nEindhoven University of Technology - Netherlands",
    "When do you wish to start your grad studies?": "2025",
    "GRE/GMAT (tick)": "-",
    "TOEFL/IELTS (tick)": "Year: 2024 Score: 8.0",
    "Current CGPA/Percentage:": "8.32",
    "LoR’s by:": "Dr. Usha BA - Professor, Department of Computer Science and Engineering",
    "Research/Work experience/Internship:": "Internship: Altsted India Private Limited - 6 months\nWork Experience: Oracle Financial Services Software Limited - 2+ years",
    "Budget, including possible loans (indicate range):": "₹ 30 Lac",
}

# Fill in form fields
for field, answer in fields.items():
    p = document.add_paragraph()
    p.add_run(field).bold = True
    p.add_run(" " + answer)

# Add signature line
document.add_paragraph("\n\nSignature of student:").add_run(" Prajwal").bold = True
document.add_paragraph("Signature of HEF").alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Save the document
document.save("HEF_Form.docx")